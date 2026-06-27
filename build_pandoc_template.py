from __future__ import annotations

import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "template.docx"
SAMPLE_IMAGE = ROOT / "verify_assets" / "template-sample.png"


COLORS = {
    "ink": "1F2328",
    "muted": "687076",
    "blue": "1F4E79",
    "blue_light": "E8EEF5",
    "gray_light": "F6F8FA",
    "gray_border": "C9D1D9",
    "quote_border": "9FB6D8",
}


def rgb(hex_color: str) -> RGBColor:
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def remove_children(element, tag: str) -> None:
    for child in list(element):
        if child.tag == qn(tag):
            element.remove(child)


def get_or_add(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_rfonts(r_pr, latin: str, east_asia: str) -> None:
    r_fonts = get_or_add(r_pr, "w:rFonts")
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:cs"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_style_font(
    style,
    *,
    latin: str = "Times New Roman",
    east_asia: str = "SimSun",
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    font = style.font
    font.name = latin
    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if color:
        font.color.rgb = rgb(color)
    set_rfonts(style.element.get_or_add_rPr(), latin, east_asia)


def set_paragraph_format(
    style,
    *,
    before: float = 0,
    after: float = 6,
    line: float = 1.25,
    first_indent_pt: float | None = None,
    left_indent_cm: float | None = None,
    right_indent_cm: float | None = None,
    alignment=None,
    keep_with_next: bool | None = None,
    page_break_before: bool | None = None,
) -> None:
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_indent_pt is not None:
        fmt.first_line_indent = Pt(first_indent_pt)
    if left_indent_cm is not None:
        fmt.left_indent = Cm(left_indent_cm)
    if right_indent_cm is not None:
        fmt.right_indent = Cm(right_indent_cm)
    if alignment is not None:
        fmt.alignment = alignment
    if keep_with_next is not None:
        fmt.keep_with_next = keep_with_next
    if page_break_before is not None:
        fmt.page_break_before = page_break_before


def set_style_shading(style, fill: str) -> None:
    p_pr = style.element.get_or_add_pPr()
    shd = get_or_add(p_pr, "w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_style_border_left(style, color: str, size: str = "8", space: str = "8") -> None:
    p_pr = style.element.get_or_add_pPr()
    p_bdr = get_or_add(p_pr, "w:pBdr")
    left = get_or_add(p_bdr, "w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)


def set_doc_defaults(doc: Document) -> None:
    styles = doc.styles.element
    doc_defaults = get_or_add(styles, "w:docDefaults")
    r_pr_default = get_or_add(doc_defaults, "w:rPrDefault")
    r_pr = get_or_add(r_pr_default, "w:rPr")
    set_rfonts(r_pr, "Times New Roman", "SimSun")
    sz = get_or_add(r_pr, "w:sz")
    sz.set(qn("w:val"), "24")
    sz_cs = get_or_add(r_pr, "w:szCs")
    sz_cs.set(qn("w:val"), "24")


def ensure_style(doc: Document, name: str, style_type):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def configure_table_style(doc: Document) -> None:
    style = ensure_style(doc, "Table", WD_STYLE_TYPE.TABLE)
    style.element.set(qn("w:customStyle"), "1")
    tbl_pr = get_or_add(style.element, "w:tblPr")

    tbl_w = get_or_add(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")

    tbl_ind = get_or_add(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    borders = get_or_add(tbl_pr, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        bdr = get_or_add(borders, f"w:{edge}")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "4")
        bdr.set(qn("w:space"), "0")
        bdr.set(qn("w:color"), COLORS["gray_border"])

    cell_mar = get_or_add(tbl_pr, "w:tblCellMar")
    for side, val in {"top": "80", "bottom": "80", "start": "120", "end": "120"}.items():
        margin = get_or_add(cell_mar, f"w:{side}")
        margin.set(qn("w:w"), val)
        margin.set(qn("w:type"), "dxa")

    remove_children(style.element, "w:tblStylePr")
    first_row = OxmlElement("w:tblStylePr")
    first_row.set(qn("w:type"), "firstRow")
    tr_pr = OxmlElement("w:trPr")
    first_row.append(tr_pr)
    tc_pr = OxmlElement("w:tcPr")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), COLORS["blue_light"])
    tc_pr.append(shd)
    first_row.append(tc_pr)
    r_pr = OxmlElement("w:rPr")
    b = OxmlElement("w:b")
    r_pr.append(b)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), COLORS["blue"])
    r_pr.append(color)
    set_rfonts(r_pr, "Times New Roman", "SimSun")
    first_row.append(r_pr)
    style.element.append(first_row)


def configure_styles(doc: Document) -> None:
    set_doc_defaults(doc)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=12, color=COLORS["ink"])
    set_paragraph_format(normal, after=6, line=1.5, first_indent_pt=24)

    for name in ("Body Text", "First Paragraph"):
        st = ensure_style(doc, name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = normal
        set_style_font(st, size=12, color=COLORS["ink"])
        set_paragraph_format(st, after=6, line=1.5, first_indent_pt=24)

    compact = ensure_style(doc, "Compact", WD_STYLE_TYPE.PARAGRAPH)
    compact.base_style = normal
    set_style_font(compact, size=12, color=COLORS["ink"])
    set_paragraph_format(compact, after=0, line=1.25, first_indent_pt=0)

    title = ensure_style(doc, "Title", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(title, latin="Times New Roman", east_asia="SimHei", size=22, bold=True, color=COLORS["ink"])
    set_paragraph_format(title, before=12, after=10, line=1.25, first_indent_pt=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    subtitle = ensure_style(doc, "Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(subtitle, east_asia="KaiTi", size=14, color=COLORS["muted"])
    set_paragraph_format(subtitle, after=12, line=1.25, first_indent_pt=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for name in ("Author", "Date"):
        st = ensure_style(doc, name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(st, east_asia="KaiTi", size=11, color=COLORS["muted"])
        set_paragraph_format(st, after=3, line=1.2, first_indent_pt=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    abstract_title = ensure_style(doc, "Abstract Title", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(abstract_title, latin="Arial", east_asia="SimHei", size=16, bold=False, color=COLORS["ink"])
    set_paragraph_format(
        abstract_title,
        before=24,
        after=18,
        line=1.0,
        first_indent_pt=0,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        page_break_before=True,
    )

    abstract = ensure_style(doc, "Abstract", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(abstract, size=12, color=COLORS["ink"])
    set_paragraph_format(abstract, after=0, line=1.25, first_indent_pt=24, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading_specs = {
        "Heading 1": (16, "SimHei", COLORS["ink"], 24, 18, WD_ALIGN_PARAGRAPH.CENTER, True),
        "Heading 2": (14, "SimHei", COLORS["ink"], 24, 6, WD_ALIGN_PARAGRAPH.LEFT, False),
        "Heading 3": (13, "SimHei", COLORS["ink"], 12, 6, WD_ALIGN_PARAGRAPH.LEFT, False),
        "Heading 4": (12, "SimHei", COLORS["ink"], 8, 4, WD_ALIGN_PARAGRAPH.LEFT, False),
        "Heading 5": (12, "SimHei", COLORS["ink"], 6, 3, WD_ALIGN_PARAGRAPH.LEFT, False),
        "Heading 6": (12, "SimHei", COLORS["ink"], 6, 3, WD_ALIGN_PARAGRAPH.LEFT, False),
    }
    for name, (size, east_asia, color, before, after, align, page_break) in heading_specs.items():
        st = ensure_style(doc, name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(st, east_asia=east_asia, size=size, bold=True, color=color)
        set_paragraph_format(
            st,
            before=before,
            after=after,
            line=1.0 if name == "Heading 1" else 1.25,
            first_indent_pt=0,
            alignment=align,
            keep_with_next=True,
            page_break_before=page_break,
        )

    block = ensure_style(doc, "Block Text", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(block, east_asia="KaiTi", size=11.5, color="2F3A45")
    set_paragraph_format(block, before=4, after=8, line=1.35, first_indent_pt=0, left_indent_cm=0.55, right_indent_cm=0.25)
    set_style_shading(block, COLORS["gray_light"])
    set_style_border_left(block, COLORS["quote_border"])

    source = ensure_style(doc, "Source Code", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(source, latin="Consolas", east_asia="DengXian", size=9.5, color="24292F")
    set_paragraph_format(source, before=3, after=3, line=1.15, first_indent_pt=0, left_indent_cm=0.25)
    set_style_shading(source, COLORS["gray_light"])

    for name in ("Caption", "Table Caption", "Image Caption"):
        st = ensure_style(doc, name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(st, east_asia="KaiTi", size=10.5, color=COLORS["muted"])
        set_paragraph_format(st, before=4, after=6, line=1.15, first_indent_pt=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for name in ("Definition Term", "Definition"):
        st = ensure_style(doc, name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(st, size=11.5, color=COLORS["ink"], bold=(name == "Definition Term"))
        set_paragraph_format(st, after=4, line=1.25, first_indent_pt=0, left_indent_cm=0.45 if name == "Definition" else 0)

    foot = ensure_style(doc, "Footnote Text", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(foot, size=9, color=COLORS["ink"])
    set_paragraph_format(foot, after=2, line=1.0, first_indent_pt=0)

    bibliography = ensure_style(doc, "Bibliography", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(bibliography, size=10.5, color=COLORS["ink"])
    set_paragraph_format(bibliography, after=3, line=1.15, first_indent_pt=-18, left_indent_cm=0.63)

    for name in ("Hyperlink", "Verbatim Char", "Footnote Reference", "Section Number", "Strong"):
        style_type = WD_STYLE_TYPE.CHARACTER
        st = ensure_style(doc, name, style_type)
        if name == "Hyperlink":
            set_style_font(st, size=12, color="0969DA")
            st.font.underline = True
        elif name == "Verbatim Char":
            set_style_font(st, latin="Consolas", east_asia="DengXian", size=10, color="24292F")
        elif name == "Strong":
            set_style_font(st, size=12, bold=True, color=COLORS["ink"])
        elif name == "Section Number":
            # Fallback for users who still pass --number-sections to Pandoc.
            # The preferred path is native Word/WPS heading numbering below.
            set_style_font(st, east_asia="SimHei", size=12, bold=True, color=COLORS["blue"])
        else:
            set_style_font(st, size=9, color=COLORS["blue"])

    toc_heading = ensure_style(doc, "TOC Heading", WD_STYLE_TYPE.PARAGRAPH)
    toc_heading.base_style = normal
    set_style_font(toc_heading, latin="Arial", east_asia="SimHei", size=16, bold=False, color=COLORS["ink"])
    set_paragraph_format(
        toc_heading,
        before=24,
        after=18,
        line=1.0,
        first_indent_pt=0,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        page_break_before=True,
    )

    configure_table_style(doc)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for item in (fld_begin, instr, fld_sep, text, fld_end):
        run._r.append(item)


def configure_sections(doc: Document) -> None:
    for name in ("Header", "Footer"):
        st = ensure_style(doc, name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(st, east_asia="SimSun", size=9, color=COLORS["muted"])
        set_paragraph_format(st, after=0, line=1.0, first_indent_pt=0)

    for section in doc.sections:
        section.start_type = WD_SECTION.NEW_PAGE
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)

        header = section.header
        header.is_linked_to_previous = False
        header.paragraphs[0].text = "Markdown 转 Word 模板"
        header.paragraphs[0].style = doc.styles["Header"]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        footer = section.footer
        footer.is_linked_to_previous = False
        footer.paragraphs[0]._p.clear_content()
        footer.paragraphs[0].style = doc.styles["Footer"]
        add_page_number(footer.paragraphs[0])


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_preview_content(doc: Document) -> None:
    doc.add_paragraph("Markdown 转 Word 模板", style="Title")
    doc.add_paragraph("Pandoc reference docx / 论文风格中文排版", style="Subtitle")
    doc.add_paragraph("用途：作为 pandoc --reference-doc=template.docx 的样式模板。", style="Abstract")

    doc.add_paragraph("模板说明", style="Heading 1")
    p = doc.add_paragraph(style="Body Text")
    p.add_run("本模板采用 A4 纸张、论文式页边距、中文正文 12 磅、1.5 倍行距。")
    p.add_run("Pandoc 转换时会读取本文档的样式定义，正文示例不会进入你的输出文档。")

    doc.add_paragraph("覆盖范围", style="Heading 2")
    for text in [
        "标题 1-6、正文、目录、脚注、引用块、定义列表、图片题注、表格题注。",
        "代码块与行内代码采用等宽字体和浅灰底色。",
        "表格使用细边框、浅蓝表头和较舒适的单元格内边距。",
    ]:
        doc.add_paragraph(text, style="Body Text")

    doc.add_paragraph("引用块示例", style="Heading 2")
    doc.add_paragraph("用于 Markdown 的 > blockquote。模板会使用楷体、浅灰底和左侧强调线，让长引用更容易区分。", style="Block Text")

    doc.add_paragraph("表格示例", style="Heading 2")
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Markdown 元素", "Word 样式", "说明"),
        ("# 标题", "Heading 1-6", "支持自动编号和目录"),
        ("代码块", "Source Code", "等宽字体、浅灰底色"),
        ("表格", "Table", "细边框、表头强调、舒适内边距"),
    ]
    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.space_after = Pt(0)


def create_sample_image() -> None:
    from PIL import Image, ImageDraw, ImageFont

    SAMPLE_IMAGE.parent.mkdir(exist_ok=True)
    image = Image.new("RGB", (1200, 560), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1199, 559), outline="#C9D1D9", width=3)
    draw.rectangle((0, 0, 1199, 96), fill="#E8EEF5")
    draw.text((48, 30), "Pandoc + template.docx", fill="#1F4E79")

    bars = [280, 420, 330, 500]
    labels = ["标题", "正文", "表格", "代码"]
    colors = ["#9FB6D8", "#8CC7A1", "#F0C36A", "#B8A1D9"]
    x = 115
    for label, height, color in zip(labels, bars, colors):
        draw.rectangle((x, 500 - height, x + 120, 500), fill=color)
        draw.text((x + 20, 515), label, fill="#1F2328")
        x += 230
    image.save(SAMPLE_IMAGE)


def patch_numbering(docx_path: Path) -> None:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        with zipfile.ZipFile(docx_path) as zin:
            zin.extractall(tmp_path)

        numbering = tmp_path / "word" / "numbering.xml"
        if numbering.exists():
            from lxml import etree

            tree = etree.parse(str(numbering))
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            for lvl in tree.xpath("//w:lvl", namespaces=ns):
                ilvl = int(lvl.get(qn("w:ilvl"), "0"))
                left = str(540 + ilvl * 360)
                hanging = "270"
                p_pr = lvl.find("w:pPr", namespaces=ns)
                if p_pr is None:
                    p_pr = etree.SubElement(lvl, qn("w:pPr"))
                ind = p_pr.find("w:ind", namespaces=ns)
                if ind is None:
                    ind = etree.SubElement(p_pr, qn("w:ind"))
                ind.set(qn("w:left"), left)
                ind.set(qn("w:hanging"), hanging)
                spacing = p_pr.find("w:spacing", namespaces=ns)
                if spacing is None:
                    spacing = etree.SubElement(p_pr, qn("w:spacing"))
                spacing.set(qn("w:after"), "80")
                spacing.set(qn("w:line"), "300")
                spacing.set(qn("w:lineRule"), "auto")
                tabs = p_pr.find("w:tabs", namespaces=ns)
                if tabs is None:
                    tabs = etree.SubElement(p_pr, qn("w:tabs"))
                for child in list(tabs):
                    tabs.remove(child)
                tab = etree.SubElement(tabs, qn("w:tab"))
                tab.set(qn("w:val"), "num")
                tab.set(qn("w:pos"), left)
            tree.write(str(numbering), xml_declaration=True, encoding="UTF-8", standalone=True)

        patch_heading_numbering(tmp_path)

        rebuilt = docx_path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(rebuilt, "w", zipfile.ZIP_DEFLATED) as zout:
            for file in tmp_path.rglob("*"):
                if file.is_file():
                    zout.write(file, file.relative_to(tmp_path))
        rebuilt.replace(docx_path)


def next_xml_id(tree, xpath: str, attr_name: str, ns: dict[str, str]) -> int:
    values = []
    for node in tree.xpath(xpath, namespaces=ns):
        raw = node.get(qn(attr_name))
        if raw and raw.isdigit():
            values.append(int(raw))
    return (max(values) + 1) if values else 1


def append_text(parent, tag: str, attrs: dict[str, str] | None = None):
    from lxml import etree

    node = etree.SubElement(parent, qn(tag))
    for key, value in (attrs or {}).items():
        node.set(qn(key), value)
    return node


def patch_heading_numbering(tmp_path: Path) -> None:
    from lxml import etree

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    numbering_path = tmp_path / "word" / "numbering.xml"
    styles_path = tmp_path / "word" / "styles.xml"

    numbering_tree = etree.parse(str(numbering_path))
    styles_tree = etree.parse(str(styles_path))

    abstract_id = next_xml_id(numbering_tree, "//w:abstractNum", "w:abstractNumId", ns)
    num_id = next_xml_id(numbering_tree, "//w:num", "w:numId", ns)

    abstract = etree.Element(qn("w:abstractNum"))
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    append_text(abstract, "w:nsid", {"w:val": "4D445750"})
    append_text(abstract, "w:multiLevelType", {"w:val": "hybridMultilevel"})
    append_text(abstract, "w:tmpl", {"w:val": "4D445750"})

    levels = [
        {"fmt": "decimal", "text": "第%1章  ", "style": "Heading1", "jc": "center", "left": "0", "hang": "0", "font": "SimHei", "size": "32"},
        {"fmt": "decimal", "text": "%1.%2  ", "style": "Heading2", "jc": "left", "left": "0", "hang": "0", "font": "SimHei", "size": "28"},
        {"fmt": "decimal", "text": "%1.%2.%3  ", "style": "Heading3", "jc": "left", "left": "0", "hang": "0", "font": "SimHei", "size": "26"},
        {"fmt": "decimal", "text": "(%4) ", "style": "Heading4", "jc": "left", "left": "0", "hang": "0", "font": "SimHei", "size": "24"},
        {"fmt": "decimal", "text": "%5)", "style": "Heading5", "jc": "left", "left": "0", "hang": "0", "font": "SimHei", "size": "22"},
        {"fmt": "lowerLetter", "text": "%6)", "style": "Heading6", "jc": "left", "left": "0", "hang": "0", "font": "SimHei", "size": "21"},
    ]

    for i, spec in enumerate(levels):
        lvl = append_text(abstract, "w:lvl", {"w:ilvl": str(i)})
        append_text(lvl, "w:start", {"w:val": "1"})
        append_text(lvl, "w:numFmt", {"w:val": spec["fmt"]})
        append_text(lvl, "w:pStyle", {"w:val": spec["style"]})
        append_text(lvl, "w:lvlText", {"w:val": spec["text"]})
        append_text(lvl, "w:lvlJc", {"w:val": spec["jc"]})

        p_pr = append_text(lvl, "w:pPr")
        append_text(p_pr, "w:autoSpaceDE", {"w:val": "0"})
        append_text(p_pr, "w:autoSpaceDN", {"w:val": "0"})
        tabs = append_text(p_pr, "w:tabs")
        append_text(tabs, "w:tab", {"w:val": "num", "w:pos": "360"})
        append_text(p_pr, "w:ind", {"w:left": spec["left"], "w:hanging": spec["hang"]})

        r_pr = append_text(lvl, "w:rPr")
        append_text(r_pr, "w:b")
        append_text(r_pr, "w:color", {"w:val": COLORS["blue"] if i < 2 else "24415F"})
        fonts = append_text(r_pr, "w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            fonts.set(qn(attr), spec["font"])
        append_text(r_pr, "w:sz", {"w:val": spec["size"]})
        append_text(r_pr, "w:szCs", {"w:val": spec["size"]})

    numbering_tree.getroot().append(abstract)

    num = etree.Element(qn("w:num"))
    num.set(qn("w:numId"), str(num_id))
    append_text(num, "w:abstractNumId", {"w:val": str(abstract_id)})
    numbering_tree.getroot().append(num)
    numbering_tree.write(str(numbering_path), xml_declaration=True, encoding="UTF-8", standalone=True)

    for style in styles_tree.xpath("//w:style", namespaces=ns):
        p_pr = style.find("w:pPr", namespaces=ns)
        if p_pr is None:
            continue
        for existing in p_pr.findall("w:numPr", namespaces=ns):
            p_pr.remove(existing)

    for level, style_id in enumerate([f"Heading{i}" for i in range(1, 7)]):
        style = styles_tree.xpath(f'//w:style[@w:styleId="{style_id}"]', namespaces=ns)
        if not style:
            continue
        p_pr = style[0].find("w:pPr", namespaces=ns)
        if p_pr is None:
            p_pr = etree.SubElement(style[0], qn("w:pPr"))
        for tag in ("w:autoSpaceDE", "w:autoSpaceDN"):
            for existing in p_pr.findall(tag, namespaces=ns):
                p_pr.remove(existing)
            node = etree.SubElement(p_pr, qn(tag))
            node.set(qn("w:val"), "0")
        for existing in p_pr.findall("w:numPr", namespaces=ns):
            p_pr.remove(existing)
        num_pr = etree.SubElement(p_pr, qn("w:numPr"))
        append_text(num_pr, "w:ilvl", {"w:val": str(level)})
        append_text(num_pr, "w:numId", {"w:val": str(num_id)})

    styles_tree.write(str(styles_path), xml_declaration=True, encoding="UTF-8", standalone=True)


def build_template() -> None:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as fh:
        subprocess.run(["pandoc", "--print-default-data-file", "reference.docx"], check=True, stdout=fh)
        base_path = Path(fh.name)

    try:
        doc = Document(base_path)
        configure_styles(doc)
        configure_sections(doc)
        clear_body(doc)
        add_preview_content(doc)
        doc.save(TEMPLATE)
        patch_numbering(TEMPLATE)
    finally:
        base_path.unlink(missing_ok=True)

    create_sample_image()


if __name__ == "__main__":
    if shutil.which("pandoc") is None:
        raise SystemExit("pandoc is required to build template.docx")
    build_template()
    print(f"created {TEMPLATE}")
    print(f"created {SAMPLE_IMAGE}")
